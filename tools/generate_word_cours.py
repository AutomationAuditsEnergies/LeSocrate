"""Génère les fichiers Word du cours à partir des .txt corrigés, sans les tags TTS.

Entrée :
  - courstxt/formation_jour1.txt
  - courstxt/formation_jour2.txt

Sortie :
  - cours/Formation_Jour1.docx
  - cours/Formation_Jour2.docx

Traitement :
  - Suppression des tags `[xxx]` (pause, calm, emphasis, warm and reassuring, etc.)
  - Nettoyage des espaces multiples et des espaces avant ponctuation
  - Un paragraphe Word par ligne non-vide du .txt source
  - Style : Calibri 12pt, justifié, interligne 1.3, marges 2.5cm
"""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

ROOT = Path("/Users/amelle/Desktop/SocrateReprise/LeSocrate")
SRC_DIR = ROOT / "courstxt"
OUT_DIR = ROOT / "cours"
OUT_DIR.mkdir(parents=True, exist_ok=True)

TAG_RE = re.compile(r"\[[^\]]*\]")
MULTI_SPACE_RE = re.compile(r" {2,}")
SPACE_PUNCT_RE = re.compile(r" +([,.;:!?])")


def strip_tags(text: str) -> str:
    text = TAG_RE.sub("", text)
    text = MULTI_SPACE_RE.sub(" ", text)
    text = SPACE_PUNCT_RE.sub(r"\1", text)
    return text.strip()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    pf.space_after = Pt(6)


def build_doc(src: Path, out: Path, title: str, subtitle: str) -> tuple[int, int]:
    raw = src.read_text(encoding="utf-8")
    lines = raw.split("\n")

    doc = Document()
    configure_doc(doc)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle)
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    written = 0
    skipped_empty = 0
    for line in lines:
        cleaned = strip_tags(line)
        if not cleaned:
            skipped_empty += 1
            continue
        p = doc.add_paragraph(cleaned)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        written += 1

    doc.save(out)
    return written, skipped_empty


def main() -> int:
    jobs = [
        (
            SRC_DIR / "formation_jour1.txt",
            OUT_DIR / "Formation_Jour1.docx",
            "Formation — Jour 1",
            "TP Employé Commercial (RNCP 37099) · texte du cours audio",
        ),
        (
            SRC_DIR / "formation_jour2.txt",
            OUT_DIR / "Formation_Jour2.docx",
            "Formation — Jour 2",
            "TP Employé Commercial (RNCP 37099) · texte du cours audio",
        ),
    ]

    for src, out, title, subtitle in jobs:
        written, skipped = build_doc(src, out, title, subtitle)
        print(f"  ✓ {out.relative_to(ROOT)} — {written} paragraphes ({skipped} lignes vides ignorées)")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
