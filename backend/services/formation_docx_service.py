"""
Génération d'un document Word .docx "programme de formation" par journée.

Remplace formation_pdf_service.py (LaTeX/xelatex) par une approche pur Python
via python-docx. Zéro dépendance système → marche sur Azure App Service Linux
sans install LaTeX.

Pipeline :
  1. Récupère le job formation (tp_name, RNCP, daily_programs JSON)
  2. Récupère les segments texte du dossier cours (sub_parts × 3 passes)
  3. Nettoie les tags Fish Audio ([pause], [warm], etc.) réservés au TTS
  4. Construit le document via python-docx
  5. Retourne les bytes .docx

Utilisation :
    docx_bytes, filename = build_course_docx(job_id=7, folder_id=42)
"""

import io
import re
import json
from datetime import datetime
from typing import Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

from database.db import get_db_connection
from utils.logger import get_logger

logger = get_logger(__name__)


# ─── Helpers texte ────────────────────────────────────────────────────────────

_FISHAUDIO_TAG_RE = re.compile(r"\[[^\[\]\n]{1,50}\]")


def _strip_tts_tags(text: str) -> str:
    """Retire les tags Fish Audio type `[pause]`, `[warm]`, `[long pause]`, etc."""
    if not text:
        return ""
    cleaned = _FISHAUDIO_TAG_RE.sub("", text)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = re.sub(r"\n[ \t]+", "\n", cleaned)
    return cleaned.strip()


def _split_paragraphs(text: str) -> list:
    """Split sur double-newline, strip, vire les vides."""
    if not text:
        return []
    return [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]


# ─── Accès DB ─────────────────────────────────────────────────────────────────

def _get_formation_job(job_id: int) -> dict:
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        """SELECT tp_name, rncp_code, total_hours, nb_days, daily_programs, platform_id
           FROM formation_pipeline_jobs WHERE id = ?""",
        (job_id,),
    )
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise ValueError(f"Job formation {job_id} introuvable")
    return {
        "tp_name": row[0],
        "rncp_code": row[1] or "",
        "total_hours": row[2],
        "nb_days": row[3],
        "daily_programs": json.loads(row[4] or "[]"),
        "platform_id": row[5],
    }


def _folder_position_in_platform(folder_id: int, platform_id: int) -> int:
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        """SELECT id FROM cours_folders
           WHERE platform_id = ?
           ORDER BY position ASC, id ASC""",
        (platform_id,),
    )
    rows = cursor.fetchall()
    conn.close()
    for idx, (fid,) in enumerate(rows):
        if fid == folder_id:
            return idx
    raise ValueError(f"Dossier {folder_id} introuvable pour plateforme {platform_id}")


def _get_segments_for_folder(folder_id: int, version: str = "current") -> list:
    """Charge les segments completed du folder.

    `version` :
    - "current" (défaut) : lit `text_content` (= état actuel, post-révision si appliquée)
    - "pre_review" : lit `text_content_pre_review` (= snapshot pris au moment
      de `_finalize_content_step`, avant que la révision ne touche au texte).
      Si la colonne est NULL pour un segment (segment plus ancien que la
      migration), fallback sur `text_content`.
    """
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        """SELECT j.id, j.sub_parts
           FROM content_generation_jobs j
           WHERE j.folder_id = ?""",
        (folder_id,),
    )
    row = cursor.fetchone()
    if not row:
        conn.close()
        raise ValueError(f"Aucun job de génération de contenu pour le dossier {folder_id}")

    cg_job_id, sub_parts_json = row
    sub_parts_names = json.loads(sub_parts_json or "[]")

    if version == "pre_review":
        # COALESCE : si la colonne pre_review est NULL (segment antérieur à
        # la migration), fallback sur text_content pour ne pas planter.
        text_col = "COALESCE(text_content_pre_review, text_content)"
    else:
        text_col = "text_content"

    cursor.execute(
        f"""SELECT sub_part_index, passe, {text_col}
           FROM content_generation_segments
           WHERE job_id = ? AND status = 'completed'
           ORDER BY sub_part_index ASC, passe ASC""",
        (cg_job_id,),
    )
    rows = cursor.fetchall()
    conn.close()

    grouped = {}
    for sub_idx, passe, text in rows:
        grouped.setdefault(sub_idx, []).append((passe, text or ""))

    result = []
    for idx, name in enumerate(sub_parts_names):
        passes = sorted(grouped.get(idx, []), key=lambda x: x[0])
        body = "\n\n".join(_strip_tts_tags(txt) for _, txt in passes).strip()
        result.append({"name": name, "body": body})

    return result


# ─── Construction du document .docx ──────────────────────────────────────────

_PURPLE = RGBColor(0x8B, 0x5C, 0xF6)  # accent violet cohérent avec le thème UI
_GREY_DARK = RGBColor(0x33, 0x33, 0x33)
_GREY_MUTED = RGBColor(0x64, 0x74, 0x8B)


def _set_margins(doc: Document, cm: float = 2.5) -> None:
    for section in doc.sections:
        section.top_margin = Cm(cm)
        section.bottom_margin = Cm(cm)
        section.left_margin = Cm(cm)
        section.right_margin = Cm(cm)


def _title(doc: Document, text: str, size: int = 28, color: RGBColor = _PURPLE) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _subtitle(doc: Document, text: str, size: int = 14, color: RGBColor = _GREY_MUTED) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int = 1) -> None:
    # Niveau 1 = titre de sub_part, niveau 2 = sous-section
    size_map = {1: 18, 2: 14}
    size = size_map.get(level, 12)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = _PURPLE if level == 1 else _GREY_DARK
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def _paragraph(doc: Document, text: str, italic: bool = False, color: RGBColor = None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if italic:
        run.italic = True
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)


def _metadata_block(doc: Document, kv_pairs: list) -> None:
    """Liste clé/valeur alignée — ex. RNCP : XXXXX — Durée : 70h."""
    for key, value in kv_pairs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_key = p.add_run(f"{key} : ")
        run_key.bold = True
        run_key.font.size = Pt(11)
        run_key.font.color.rgb = _GREY_DARK
        run_val = p.add_run(str(value))
        run_val.font.size = Pt(11)
        run_val.font.color.rgb = _GREY_DARK


def _horizontal_rule(doc: Document) -> None:
    """Trait de séparation — ligne de tirets centrés (python-docx n'offre pas de HR natif)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 40)
    run.font.color.rgb = _GREY_MUTED
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


# ─── API publique ─────────────────────────────────────────────────────────────

def build_course_docx(job_id: int, folder_id: int, version: str = "current") -> Tuple[bytes, str]:
    """
    Produit un document Word .docx d'une journée de formation.

    `version` :
    - "current" (défaut) : texte actuel en DB (= post-révision si appliquée)
    - "pre_review" : snapshot pris au finalize content (= avant révision)

    Retourne (docx_bytes, suggested_filename).
    """
    job = _get_formation_job(job_id)
    position = _folder_position_in_platform(folder_id, job["platform_id"])
    daily_programs = job["daily_programs"]

    if position >= len(daily_programs):
        raise ValueError(
            f"Position dossier ({position}) dépasse nb_days ({len(daily_programs)}) — "
            f"le dossier n'appartient pas à la pipeline formation"
        )
    day_data = daily_programs[position]
    day_number = day_data.get("day_number", position + 1)
    day_title = day_data.get("title", f"Journée {day_number}")

    segments = _get_segments_for_folder(folder_id, version=version)
    module_by_name = {
        sp.get("name"): sp.get("module_content", "")
        for sp in day_data.get("sub_parts", [])
    }

    total_words = 0
    sub_parts_render = []
    for part in segments:
        body = part["body"]
        total_words += len(body.split())
        sub_parts_render.append({
            "name": part["name"],
            "module_content": module_by_name.get(part["name"], ""),
            "body": body,
        })

    # ─── Construction du document ───────────────────────────────────────────
    doc = Document()
    _set_margins(doc, cm=2.5)

    # Page de garde
    _title(doc, job["tp_name"])
    _subtitle(doc, f"Journée {day_number} — {day_title}")
    doc.add_paragraph()
    _metadata_block(doc, [
        ("Code RNCP", job["rncp_code"] or "—"),
        ("Durée totale", f"{job['total_hours']}h sur {job['nb_days']} journées"),
        ("Volume de cette journée", f"{total_words:,} mots".replace(",", " ")),
        ("Généré le", datetime.now().strftime("%d/%m/%Y à %H:%M")),
    ])
    _horizontal_rule(doc)

    # Résumé de la journée
    day_recap = _strip_tts_tags(day_data.get("day_recap") or "")
    if day_recap:
        _heading(doc, "Résumé de la journée", level=2)
        for para in _split_paragraphs(day_recap):
            _paragraph(doc, para, italic=True, color=_GREY_MUTED)

    # Transition pédagogique
    day_transition = _strip_tts_tags(day_data.get("day_transition") or "")
    if day_transition:
        _heading(doc, "Transition pédagogique", level=2)
        for para in _split_paragraphs(day_transition):
            _paragraph(doc, para, italic=True, color=_GREY_MUTED)

    # Les sub_parts (sections du cours)
    for idx, part in enumerate(sub_parts_render, 1):
        doc.add_page_break()
        _heading(doc, f"{idx}. {part['name']}", level=1)

        if part["module_content"]:
            _heading(doc, "Plan pédagogique", level=2)
            for para in _split_paragraphs(part["module_content"]):
                _paragraph(doc, para, italic=True, color=_GREY_MUTED)

        if part["body"]:
            _heading(doc, "Contenu du cours", level=2)
            for para in _split_paragraphs(part["body"]):
                _paragraph(doc, para)

    # Sérialisation
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    slug = re.sub(r"[^a-zA-Z0-9]+", "-", job["tp_name"].lower()).strip("-")[:40]
    suffix = "" if version == "current" else f"-{version.replace('_', '-')}"
    filename = f"programme-{slug}-jour{day_number}{suffix}.docx"
    logger.info(f"✅ DOCX généré : {filename} ({len(docx_bytes):,} bytes, {total_words} mots, version={version})")
    return docx_bytes, filename
